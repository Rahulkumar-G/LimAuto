"""
Advanced Export and Publishing Pipeline
Professional-grade document generation and publishing workflow
"""

from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional
import json
import os
import asyncio
from pathlib import Path
import zipfile
import io

from .app import db, User, BookProject, celery_app, redis_client
from .subscription_manager import subscription_manager
from ..utils.logger import get_logger

export_bp = Blueprint('export', __name__, url_prefix='/api/export')
logger = get_logger(__name__)

class ExportEngine:
    """Advanced export engine with multiple format support"""
    
    def __init__(self):
        self.redis = redis_client
        self.logger = logger
        self.export_formats = {
            'pdf': {
                'name': 'PDF Document',
                'description': 'Professional PDF with custom styling',
                'file_extension': '.pdf',
                'mime_type': 'application/pdf',
                'features': ['Custom fonts', 'Headers/footers', 'Table of contents', 'Bookmarks']
            },
            'epub': {
                'name': 'EPUB E-book',
                'description': 'E-reader compatible format',
                'file_extension': '.epub',
                'mime_type': 'application/epub+zip',
                'features': ['Responsive layout', 'Metadata', 'Chapter navigation', 'Images']
            },
            'docx': {
                'name': 'Microsoft Word',
                'description': 'Editable Word document',
                'file_extension': '.docx',
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'features': ['Styles', 'Comments', 'Track changes', 'Headers/footers']
            },
            'html': {
                'name': 'HTML Website',
                'description': 'Web-ready HTML with CSS',
                'file_extension': '.html',
                'mime_type': 'text/html',
                'features': ['Responsive design', 'Search functionality', 'Print styles', 'Analytics']
            },
            'latex': {
                'name': 'LaTeX Source',
                'description': 'Academic publication ready',
                'file_extension': '.tex',
                'mime_type': 'application/x-tex',
                'features': ['Bibliography', 'Citations', 'Mathematical formulas', 'Professional typography']
            },
            'markdown': {
                'name': 'Markdown',
                'description': 'Plain text with formatting',
                'file_extension': '.md',
                'mime_type': 'text/markdown',
                'features': ['Version control friendly', 'GitHub compatible', 'Lightweight', 'Extensible']
            }
        }
    
    async def generate_export(self, project_id: str, format_type: str, 
                            options: Dict[str, Any] = None) -> Dict[str, Any]:
        """Generate export in specified format"""
        try:
            project = BookProject.query.get(project_id)
            if not project:
                return {'error': 'Project not found'}
            
            if format_type not in self.export_formats:
                return {'error': f'Unsupported format: {format_type}'}
            
            options = options or {}
            
            # Get project content
            content_data = self._prepare_content_data(project)
            
            # Format-specific generation
            if format_type == 'pdf':
                result = await self._generate_pdf(project, content_data, options)
            elif format_type == 'epub':
                result = await self._generate_epub(project, content_data, options)
            elif format_type == 'docx':
                result = await self._generate_docx(project, content_data, options)
            elif format_type == 'html':
                result = await self._generate_html(project, content_data, options)
            elif format_type == 'latex':
                result = await self._generate_latex(project, content_data, options)
            elif format_type == 'markdown':
                result = await self._generate_markdown(project, content_data, options)
            else:
                return {'error': 'Format not implemented'}
            
            return result
            
        except Exception as e:
            self.logger.error(f"Export generation error: {e}")
            return {'error': str(e)}
    
    def _prepare_content_data(self, project: BookProject) -> Dict[str, Any]:
        """Prepare content data for export"""
        outline = json.loads(project.outline_json) if project.outline_json else []
        chapters = json.loads(project.chapters_json) if project.chapters_json else {}
        metadata = json.loads(project.metadata_json) if project.metadata_json else {}
        
        return {
            'project': {
                'id': project.id,
                'title': project.title,
                'topic': project.topic,
                'description': project.description,
                'author': project.author.full_name,
                'created_at': project.created_at,
                'updated_at': project.updated_at,
                'word_count': project.word_count,
                'quality_score': project.quality_score
            },
            'outline': outline,
            'chapters': chapters,
            'metadata': metadata
        }
    
    async def _generate_pdf(self, project: BookProject, content_data: Dict, 
                          options: Dict) -> Dict[str, Any]:
        """Generate professional PDF with LaTeX/ReportLab"""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        import markdown
        
        # PDF options
        page_size = A4 if options.get('page_size') == 'A4' else letter
        include_toc = options.get('include_toc', True)
        include_headers = options.get('include_headers', True)
        custom_styling = options.get('custom_styling', {})
        
        # Create PDF buffer
        buffer = io.BytesIO()
        
        # Create document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=page_size,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#2C3E50')
        )
        
        chapter_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            textColor=colors.HexColor('#34495E')
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            leftIndent=0,
            rightIndent=0
        )
        
        # Story elements
        story = []
        
        # Title page
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(project.title, title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"by {project.author.full_name}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(PageBreak())
        
        # Table of contents
        if include_toc and content_data['outline']:
            story.append(Paragraph("Table of Contents", styles['Heading1']))
            story.append(Spacer(1, 20))
            
            toc_data = []
            for i, chapter_title in enumerate(content_data['outline'], 1):
                toc_data.append([f"Chapter {i}", chapter_title, f"{i*10}"])
            
            toc_table = Table(toc_data, colWidths=[1*inch, 4*inch, 1*inch])
            toc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ]))
            story.append(toc_table)
            story.append(PageBreak())
        
        # Chapters
        for i, chapter_title in enumerate(content_data['outline'], 1):
            # Chapter title
            story.append(Paragraph(f"Chapter {i}: {chapter_title}", chapter_style))
            story.append(Spacer(1, 20))
            
            # Chapter content
            chapter_content = content_data['chapters'].get(chapter_title, '')
            if chapter_content:
                # Convert markdown to HTML then to ReportLab
                html_content = markdown.markdown(chapter_content)
                # Simplified HTML to ReportLab conversion
                paragraphs = html_content.split('</p>')
                for para in paragraphs:
                    clean_para = para.replace('<p>', '').replace('<strong>', '<b>').replace('</strong>', '</b>')
                    if clean_para.strip():
                        story.append(Paragraph(clean_para, body_style))
                        story.append(Spacer(1, 12))
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Save to file
        export_dir = Path('exports') / project.id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"{project.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        filepath = export_dir / filename
        
        with open(filepath, 'wb') as f:
            f.write(buffer.getvalue())
        
        return {
            'format': 'pdf',
            'filename': filename,
            'filepath': str(filepath),
            'file_size': len(buffer.getvalue()),
            'pages_estimated': len(content_data['outline']) + 2,
            'generated_at': datetime.utcnow().isoformat()
        }
    
    async def _generate_epub(self, project: BookProject, content_data: Dict, 
                           options: Dict) -> Dict[str, Any]:
        """Generate EPUB e-book format"""
        from ebooklib import epub
        import markdown
        
        # Create EPUB book
        book = epub.EpubBook()
        
        # Metadata
        book.set_identifier(f"booklm_{project.id}")
        book.set_title(project.title)
        book.set_language('en')
        book.add_author(project.author.full_name)
        book.add_metadata('DC', 'description', project.description or project.topic)
        book.add_metadata('DC', 'subject', project.topic)
        
        # Style
        default_css = epub.EpubItem(
            uid="default",
            file_name="style/default.css",
            media_type="text/css",
            content="""
            body { font-family: Georgia, serif; margin: 1em; }
            h1 { color: #2C3E50; border-bottom: 2px solid #3498DB; padding-bottom: 0.3em; }
            h2 { color: #34495E; margin-top: 2em; }
            p { text-align: justify; line-height: 1.6; margin-bottom: 1em; }
            .chapter { page-break-before: always; }
            """
        )
        book.add_item(default_css)
        
        # Chapters
        chapters = []
        for i, chapter_title in enumerate(content_data['outline'], 1):
            chapter_content = content_data['chapters'].get(chapter_title, '')
            
            if chapter_content:
                # Convert markdown to HTML
                html_content = markdown.markdown(chapter_content)
                
                # Create EPUB chapter
                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=f'chapter_{i}.xhtml',
                    lang='en'
                )
                
                chapter.content = f"""
                <html xmlns="http://www.w3.org/1999/xhtml">
                <head>
                    <title>{chapter_title}</title>
                    <link rel="stylesheet" href="style/default.css" type="text/css"/>
                </head>
                <body>
                    <div class="chapter">
                        <h1>Chapter {i}: {chapter_title}</h1>
                        {html_content}
                    </div>
                </body>
                </html>
                """
                
                book.add_item(chapter)
                chapters.append(chapter)
        
        # Table of contents
        book.toc = [(epub.Section('Chapters'), chapters)]
        
        # Navigation
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Spine
        book.spine = ['nav'] + chapters
        
        # Save EPUB
        export_dir = Path('exports') / project.id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"{project.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.epub"
        filepath = export_dir / filename
        
        epub.write_epub(str(filepath), book)
        
        file_size = filepath.stat().st_size
        
        return {
            'format': 'epub',
            'filename': filename,
            'filepath': str(filepath),
            'file_size': file_size,
            'chapters': len(chapters),
            'generated_at': datetime.utcnow().isoformat()
        }
    
    async def _generate_docx(self, project: BookProject, content_data: Dict, 
                           options: Dict) -> Dict[str, Any]:
        """Generate Microsoft Word document"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import markdown
        from bs4 import BeautifulSoup
        
        # Create document
        doc = Document()
        
        # Document properties
        doc.core_properties.title = project.title
        doc.core_properties.author = project.author.full_name
        doc.core_properties.subject = project.topic
        doc.core_properties.comments = project.description
        
        # Styles
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(30)
        
        # Chapter style
        chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
        chapter_style.font.name = 'Calibri'
        chapter_style.font.size = Pt(18)
        chapter_style.font.bold = True
        chapter_style.paragraph_format.space_before = Pt(24)
        chapter_style.paragraph_format.space_after = Pt(12)
        
        # Title page
        title_para = doc.add_paragraph(project.title, style='BookTitle')
        
        author_para = doc.add_paragraph(f"by {project.author.full_name}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of contents
        if options.get('include_toc', True):
            toc_para = doc.add_paragraph('Table of Contents')
            toc_para.style = 'Heading 1'
            
            for i, chapter_title in enumerate(content_data['outline'], 1):
                toc_item = doc.add_paragraph(f"Chapter {i}: {chapter_title}")
                toc_item.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_page_break()
        
        # Chapters
        for i, chapter_title in enumerate(content_data['outline'], 1):
            # Chapter title
            chapter_para = doc.add_paragraph(f"Chapter {i}: {chapter_title}", style='ChapterTitle')
            
            # Chapter content
            chapter_content = content_data['chapters'].get(chapter_title, '')
            if chapter_content:
                # Convert markdown to HTML then extract text
                html_content = markdown.markdown(chapter_content)
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Process paragraphs
                for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol']):
                    if element.name in ['h1', 'h2', 'h3']:
                        heading_para = doc.add_paragraph(element.get_text())
                        heading_para.style = f'Heading {element.name[1]}'
                    elif element.name == 'p':
                        doc.add_paragraph(element.get_text())
                    elif element.name in ['ul', 'ol']:
                        for li in element.find_all('li'):
                            list_para = doc.add_paragraph(li.get_text())
                            list_para.style = 'List Bullet' if element.name == 'ul' else 'List Number'
            
            doc.add_page_break()
        
        # Save document
        export_dir = Path('exports') / project.id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"{project.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = export_dir / filename
        
        doc.save(str(filepath))
        
        file_size = filepath.stat().st_size
        
        return {
            'format': 'docx',
            'filename': filename,
            'filepath': str(filepath),
            'file_size': file_size,
            'pages_estimated': len(content_data['outline']) * 5,
            'generated_at': datetime.utcnow().isoformat()
        }
    
    async def _generate_html(self, project: BookProject, content_data: Dict, 
                           options: Dict) -> Dict[str, Any]:
        """Generate responsive HTML website"""
        import markdown
        
        # HTML template
        html_template = """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{title}</title>
            <meta name="description" content="{description}">
            <meta name="author" content="{author}">
            <style>
                * {{ margin: 0; padding: 0; box-sizing: border-box; }}
                body {{ 
                    font-family: 'Georgia', serif; 
                    line-height: 1.6; 
                    color: #333; 
                    max-width: 800px; 
                    margin: 0 auto; 
                    padding: 20px;
                    background: #f8f9fa;
                }}
                .container {{ background: white; padding: 40px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
                h1 {{ color: #2C3E50; border-bottom: 3px solid #3498DB; padding-bottom: 10px; margin-bottom: 30px; }}
                h2 {{ color: #34495E; margin-top: 40px; margin-bottom: 20px; }}
                h3 {{ color: #7F8C8D; margin-top: 30px; margin-bottom: 15px; }}
                p {{ margin-bottom: 16px; text-align: justify; }}
                .title-page {{ text-align: center; margin-bottom: 50px; padding: 40px 0; }}
                .book-title {{ font-size: 2.5em; margin-bottom: 20px; color: #2C3E50; }}
                .author {{ font-size: 1.2em; color: #7F8C8D; margin-bottom: 10px; }}
                .date {{ color: #95A5A6; }}
                .chapter {{ margin-bottom: 50px; page-break-before: always; }}
                .toc {{ margin-bottom: 40px; }}
                .toc ul {{ list-style: none; }}
                .toc li {{ margin: 10px 0; padding-left: 20px; }}
                .toc a {{ text-decoration: none; color: #3498DB; }}
                .toc a:hover {{ text-decoration: underline; }}
                @media print {{ 
                    body {{ background: white; }}
                    .container {{ box-shadow: none; }}
                    .chapter {{ page-break-before: always; }}
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="title-page">
                    <h1 class="book-title">{title}</h1>
                    <p class="author">by {author}</p>
                    <p class="date">Generated on {date}</p>
                </div>
                
                {toc}
                
                {content}
            </div>
        </body>
        </html>
        """
        
        # Generate table of contents
        toc_html = ""
        if options.get('include_toc', True):
            toc_html = '<div class="toc"><h2>Table of Contents</h2><ul>'
            for i, chapter_title in enumerate(content_data['outline'], 1):
                chapter_id = f"chapter-{i}"
                toc_html += f'<li><a href="#{chapter_id}">Chapter {i}: {chapter_title}</a></li>'
            toc_html += '</ul></div>'
        
        # Generate content
        content_html = ""
        for i, chapter_title in enumerate(content_data['outline'], 1):
            chapter_id = f"chapter-{i}"
            content_html += f'<div class="chapter" id="{chapter_id}">'
            content_html += f'<h1>Chapter {i}: {chapter_title}</h1>'
            
            chapter_content = content_data['chapters'].get(chapter_title, '')
            if chapter_content:
                html_content = markdown.markdown(chapter_content)
                content_html += html_content
            
            content_html += '</div>'
        
        # Fill template
        final_html = html_template.format(
            title=project.title,
            description=project.description or project.topic,
            author=project.author.full_name,
            date=datetime.now().strftime('%B %d, %Y'),
            toc=toc_html,
            content=content_html
        )
        
        # Save HTML
        export_dir = Path('exports') / project.id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"{project.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.html"
        filepath = export_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(final_html)
        
        file_size = filepath.stat().st_size
        
        return {
            'format': 'html',
            'filename': filename,
            'filepath': str(filepath),
            'file_size': file_size,
            'chapters': len(content_data['outline']),
            'generated_at': datetime.utcnow().isoformat()
        }
    
    async def _generate_latex(self, project: BookProject, content_data: Dict, 
                            options: Dict) -> Dict[str, Any]:
        """Generate LaTeX source for academic publishing"""
        # LaTeX template
        latex_template = """
\\documentclass[11pt,a4paper]{{book}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[english]{{babel}}
\\usepackage{{geometry}}
\\usepackage{{fancyhdr}}
\\usepackage{{titlesec}}
\\usepackage{{tocloft}}
\\usepackage{{hyperref}}

\\geometry{{margin=1in}}

\\title{{{title}}}
\\author{{{author}}}
\\date{{{date}}}

\\begin{{document}}

\\maketitle

\\tableofcontents
\\newpage

{content}

\\end{{document}}
        """
        
        # Generate content
        content_latex = ""
        for i, chapter_title in enumerate(content_data['outline'], 1):
            content_latex += f"\\chapter{{{chapter_title}}}\\n\\n"
            
            chapter_content = content_data['chapters'].get(chapter_title, '')
            if chapter_content:
                # Simple markdown to LaTeX conversion
                latex_content = chapter_content.replace('#', '\\section')
                latex_content = latex_content.replace('**', '\\textbf')
                latex_content = latex_content.replace('*', '\\textit')
                content_latex += latex_content + "\\n\\n"
        
        # Fill template
        final_latex = latex_template.format(
            title=project.title,
            author=project.author.full_name,
            date=datetime.now().strftime('%B %d, %Y'),
            content=content_latex
        )
        
        # Save LaTeX
        export_dir = Path('exports') / project.id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"{project.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.tex"
        filepath = export_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(final_latex)
        
        file_size = filepath.stat().st_size
        
        return {
            'format': 'latex',
            'filename': filename,
            'filepath': str(filepath),
            'file_size': file_size,
            'chapters': len(content_data['outline']),
            'generated_at': datetime.utcnow().isoformat()
        }
    
    async def _generate_markdown(self, project: BookProject, content_data: Dict, 
                               options: Dict) -> Dict[str, Any]:
        """Generate clean Markdown format"""
        
        # Combine all content
        markdown_content = f"# {project.title}\\n\\n"
        markdown_content += f"**Author:** {project.author.full_name}\\n"
        markdown_content += f"**Generated:** {datetime.now().strftime('%B %d, %Y')}\\n"
        markdown_content += f"**Topic:** {project.topic}\\n\\n"
        
        if project.description:
            markdown_content += f"## Description\\n\\n{project.description}\\n\\n"
        
        # Table of contents
        if options.get('include_toc', True):
            markdown_content += "## Table of Contents\\n\\n"
            for i, chapter_title in enumerate(content_data['outline'], 1):
                chapter_anchor = chapter_title.lower().replace(' ', '-').replace(',', '').replace('.', '')
                markdown_content += f"{i}. [{chapter_title}](#{chapter_anchor})\\n"
            markdown_content += "\\n"
        
        # Chapters
        for i, chapter_title in enumerate(content_data['outline'], 1):
            markdown_content += f"## Chapter {i}: {chapter_title}\\n\\n"
            
            chapter_content = content_data['chapters'].get(chapter_title, '')
            if chapter_content:
                markdown_content += chapter_content + "\\n\\n"
            
            markdown_content += "---\\n\\n"
        
        # Save Markdown
        export_dir = Path('exports') / project.id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"{project.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.md"
        filepath = export_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        file_size = filepath.stat().st_size
        
        return {
            'format': 'markdown',
            'filename': filename,
            'filepath': str(filepath),
            'file_size': file_size,
            'chapters': len(content_data['outline']),
            'generated_at': datetime.utcnow().isoformat()
        }

# Initialize export engine
export_engine = ExportEngine()

@export_bp.route('/formats', methods=['GET'])
@jwt_required()
def get_export_formats():
    """Get available export formats based on subscription"""
    try:
        user_id = get_jwt_identity()
        user = User.query.get(user_id)
        
        # Get user's available formats based on subscription
        subscription_info = subscription_manager.get_user_limits(user)
        available_formats = subscription_info['plan']['features']['export_formats']
        
        # Filter formats
        user_formats = {}
        for format_id, format_info in export_engine.export_formats.items():
            if format_id in available_formats:
                user_formats[format_id] = {
                    **format_info,
                    'available': True
                }
            else:
                user_formats[format_id] = {
                    **format_info,
                    'available': False,
                    'upgrade_required': True
                }
        
        return jsonify({
            'formats': user_formats,
            'subscription_tier': user.subscription_tier,
            'total_available': len(available_formats)
        }), 200
        
    except Exception as e:
        logger.error(f"Get export formats error: {e}")
        return jsonify({'error': 'Failed to load export formats'}), 500

@export_bp.route('/generate', methods=['POST'])
@jwt_required()
def generate_export():
    """Generate export in specified format"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        
        project_id = data.get('project_id')
        format_type = data.get('format')
        options = data.get('options', {})
        
        if not project_id or not format_type:
            return jsonify({'error': 'Project ID and format are required'}), 400
        
        # Check project access
        project = BookProject.query.filter(
            (BookProject.id == project_id) & 
            ((BookProject.user_id == user_id) | 
             (BookProject.collaborators.any(user_id=user_id)))
        ).first()
        
        if not project:
            return jsonify({'error': 'Project not found or access denied'}), 404
        
        # Check subscription and format availability
        user = User.query.get(user_id)
        can_export, message = subscription_manager.can_perform_action(
            user, 'export_format', format=format_type
        )
        if not can_export:
            return jsonify({'error': message}), 402
        
        # Start background export task
        task = generate_export_async.delay(project_id, format_type, options, user_id)
        
        return jsonify({
            'message': 'Export generation started',
            'task_id': task.id,
            'format': format_type,
            'estimated_time': '2-5 minutes',
            'status_endpoint': f'/api/export/status/{task.id}'
        }), 202
        
    except Exception as e:
        logger.error(f"Generate export error: {e}")
        return jsonify({'error': 'Failed to start export generation'}), 500

@export_bp.route('/status/<task_id>', methods=['GET'])
@jwt_required()
def get_export_status(task_id: str):
    """Get export generation status"""
    try:
        task = generate_export_async.AsyncResult(task_id)
        
        if task.state == 'PENDING':
            response = {
                'state': 'PENDING',
                'status': 'Export is queued for processing'
            }
        elif task.state == 'PROGRESS':
            response = {
                'state': 'PROGRESS',
                'status': task.info.get('status', ''),
                'progress': task.info.get('progress', 0)
            }
        elif task.state == 'SUCCESS':
            response = {
                'state': 'SUCCESS',
                'result': task.result
            }
        else:  # FAILURE
            response = {
                'state': 'FAILURE',
                'error': str(task.info)
            }
        
        return jsonify(response), 200
        
    except Exception as e:
        logger.error(f"Export status error: {e}")
        return jsonify({'error': 'Failed to get export status'}), 500

@export_bp.route('/download/<task_id>', methods=['GET'])
@jwt_required()
def download_export(task_id: str):
    """Download generated export file"""
    try:
        user_id = get_jwt_identity()
        
        # Get task result
        task = generate_export_async.AsyncResult(task_id)
        
        if task.state != 'SUCCESS':
            return jsonify({'error': 'Export not ready or failed'}), 400
        
        result = task.result
        filepath = result.get('filepath')
        
        if not filepath or not os.path.exists(filepath):
            return jsonify({'error': 'Export file not found'}), 404
        
        # Verify user has access to this export
        # (Additional security check would be implemented here)
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=result.get('filename'),
            mimetype=export_engine.export_formats[result['format']]['mime_type']
        )
        
    except Exception as e:
        logger.error(f"Download export error: {e}")
        return jsonify({'error': 'Failed to download export'}), 500

@celery_app.task(bind=True)
def generate_export_async(self, project_id: str, format_type: str, 
                         options: Dict, user_id: str):
    """Asynchronous export generation task"""
    try:
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'status': 'Starting export generation...', 'progress': 10}
        )
        
        # Generate export
        result = asyncio.run(
            export_engine.generate_export(project_id, format_type, options)
        )
        
        if 'error' in result:
            raise Exception(result['error'])
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'status': 'Finalizing export...', 'progress': 90}
        )
        
        # Track usage
        user = User.query.get(user_id)
        subscription_manager.track_usage(user, 'export_generated', format=format_type)
        
        return result
        
    except Exception as e:
        logger.error(f"Export generation task error: {e}")
        raise

@export_bp.route('/bulk', methods=['POST'])
@jwt_required()
def bulk_export():
    """Generate multiple export formats simultaneously"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        
        project_id = data.get('project_id')
        formats = data.get('formats', [])
        options = data.get('options', {})
        
        if not project_id or not formats:
            return jsonify({'error': 'Project ID and formats are required'}), 400
        
        # Check project access
        project = BookProject.query.filter(
            (BookProject.id == project_id) & 
            ((BookProject.user_id == user_id) | 
             (BookProject.collaborators.any(user_id=user_id)))
        ).first()
        
        if not project:
            return jsonify({'error': 'Project not found or access denied'}), 404
        
        # Check subscription
        user = User.query.get(user_id)
        can_bulk, message = subscription_manager.can_perform_action(user, 'bulk_generation')
        if not can_bulk:
            return jsonify({'error': message}), 402
        
        # Start bulk export tasks
        tasks = {}
        for format_type in formats:
            can_export, format_message = subscription_manager.can_perform_action(
                user, 'export_format', format=format_type
            )
            if can_export:
                task = generate_export_async.delay(project_id, format_type, options, user_id)
                tasks[format_type] = task.id
        
        return jsonify({
            'message': 'Bulk export generation started',
            'tasks': tasks,
            'total_formats': len(tasks),
            'estimated_time': '5-15 minutes'
        }), 202
        
    except Exception as e:
        logger.error(f"Bulk export error: {e}")
        return jsonify({'error': 'Failed to start bulk export'}), 500